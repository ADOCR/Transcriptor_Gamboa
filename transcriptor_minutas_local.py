# -*- coding: utf-8 -*-
"""
Transcriptor local para Windows/Jupyter/Anaconda.

Uso recomendado en Jupyter:
1. Ejecuta la celda de instalacion:
   %pip install -U faster-whisper python-docx tqdm requests
2. Pega este archivo completo en una celda.
3. Ejecuta:
   resultados = main()

Requisito externo:
- Instalar Ollama desde https://ollama.com
- Descargar modelo local:
  ollama pull qwen3:8b

Opcional:
- conda install -c conda-forge ffmpeg

Notas de VRAM para RTX 3070 8 GB:
- Para maxima calidad: WHISPER_MODEL_NAME = "large-v3"
- Para mas velocidad/menos VRAM: WHISPER_MODEL_NAME = "medium"
- Por estabilidad en Jupyter/Windows, la limpieza CUDA agresiva esta apagada.
- Si haces transcripcion + minuta en una sola corrida, se transcribe primero y
  despues se llama a Ollama.
- Si prefieres separar trabajo, primero usa "Solo transcripcion" y despues
  "Generar resumen/minuta desde TXT existente".
"""

from __future__ import annotations

import gc
import re
import textwrap
import traceback
from pathlib import Path
from typing import Dict, Iterable, List, Optional, Sequence, Tuple, Union

import requests
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from tqdm import tqdm


# =============================================================================
# CONFIGURACION GENERAL
# =============================================================================

# Whisper
WHISPER_MODEL_NAME = "large-v3"
LANGUAGE = "es"
TASK = "transcribe"
BEAM_SIZE = 5
VAD_FILTER = True

# GPU CUDA. Si falla, el script cae automaticamente a CPU.
GPU_DEVICE = "cuda"
GPU_COMPUTE_TYPE = "float16"

# CPU de respaldo.
CPU_DEVICE = "cpu"
CPU_COMPUTE_TYPE = "int8"
CPU_THREADS = 6

# Ollama local
GENERAR_MINUTA = True
OLLAMA_MODEL = "qwen3:8b"
OLLAMA_TIMEOUT = 900

OLLAMA_OPTIONS = {
    "temperature": 0.1,
    # Para RTX 3070 8 GB, 4096 suele ser mas estable si haces todo en una corrida.
    # Si generas la minuta por separado y tienes VRAM libre, puedes probar:
    # "num_ctx": 8192,
    "num_ctx": 4096,
}

# Con num_ctx=4096 conviene usar fragmentos moderados.
# Si usas num_ctx=8192 y el modelo responde bien, puedes subir esto a 12000.
TAMANO_FRAGMENTO_MINUTA = 8000

# Salidas
CARPETA_SALIDA = "transcripciones"
SOBRESCRIBIR_SALIDAS = False

# Memoria/estabilidad en Jupyter:
# False = mas estable; no fuerza liberacion de VRAM antes de Ollama.
# True = intenta soltar la referencia de Whisper antes de generar minutas.
LIBERAR_WHISPER_ANTES_OLLAMA = False

# Si no se libera Whisper, se conserva una referencia global para que Python no
# intente destruir el modelo al final de la funcion. Esto prioriza estabilidad
# del kernel sobre liberar VRAM inmediatamente.
RETENER_MODELO_WHISPER_EN_MEMORIA = True

# Mantener False salvo que sepas que tu entorno lo tolera.
# En algunos Jupyter + Windows + CUDA estas limpiezas pueden tumbar el kernel.
FORZAR_GC_AL_LIBERAR = False
LIMPIEZA_CUDA_AGRESIVA = False

_MODELO_WHISPER_RETENIDO = None
_MODELO_WHISPER_RETENIDO_INFO = None
_ULTIMO_WHISPER_INFO = None


# =============================================================================
# FORMATOS SOPORTADOS
# =============================================================================

EXTENSIONES_AUDIO = {
    ".mp3",
    ".wav",
    ".m4a",
    ".aac",
    ".flac",
    ".ogg",
    ".wma",
}

EXTENSIONES_VIDEO = {
    ".mp4",
    ".mkv",
    ".mov",
    ".avi",
    ".webm",
    ".m4v",
    ".wmv",
}

EXTENSIONES_COMPATIBLES = EXTENSIONES_AUDIO | EXTENSIONES_VIDEO

NOMBRES_RESERVADOS_WINDOWS = {
    "CON",
    "PRN",
    "AUX",
    "NUL",
    *(f"COM{i}" for i in range(1, 10)),
    *(f"LPT{i}" for i in range(1, 10)),
}


# =============================================================================
# UTILIDADES DE INTERFAZ Y RUTAS
# =============================================================================

def elegir_modo_proceso() -> int:
    """Pregunta el modo principal de trabajo."""
    print("\nQue deseas hacer?")
    print("1 = Solo transcripcion")
    print("2 = Transcripcion + resumen/minuta")
    print("3 = Generar resumen/minuta desde una transcripcion TXT existente")

    while True:
        opcion = input("Elige 1, 2 o 3: ").strip()
        if opcion in {"1", "2", "3"}:
            return int(opcion)
        print("Opcion no valida. Escribe 1, 2 o 3.")


def _crear_ventana_tk():
    import tkinter as tk

    root = tk.Tk()
    root.withdraw()
    root.attributes("-topmost", True)
    root.update()
    return root


def elegir_archivo_o_carpeta() -> Tuple[str, Optional[Path]]:
    """Permite elegir un archivo individual o una carpeta completa."""
    from tkinter import filedialog

    print("\nQue deseas seleccionar?")
    print("1 = Archivo individual")
    print("2 = Carpeta completa")

    while True:
        opcion = input("Elige 1 o 2: ").strip()
        if opcion in {"1", "2"}:
            break
        print("Opcion no valida. Escribe 1 o 2.")

    root = _crear_ventana_tk()
    try:
        if opcion == "1":
            patrones = " ".join(f"*{ext}" for ext in sorted(EXTENSIONES_COMPATIBLES))
            ruta = filedialog.askopenfilename(
                title="Selecciona un archivo de audio o video",
                filetypes=[
                    ("Audio y video compatibles", patrones),
                    ("Todos los archivos", "*.*"),
                ],
            )
            return "archivo", Path(ruta) if ruta else None

        ruta = filedialog.askdirectory(title="Selecciona una carpeta")
        return "carpeta", Path(ruta) if ruta else None
    finally:
        root.destroy()


def seleccionar_txt_limpio() -> Optional[Path]:
    """Abre selector para elegir una transcripcion TXT existente."""
    from tkinter import filedialog

    root = _crear_ventana_tk()
    try:
        ruta = filedialog.askopenfilename(
            title="Selecciona un TXT limpio para generar minuta",
            filetypes=[
                ("Archivos TXT", "*.txt"),
                ("Todos los archivos", "*.*"),
            ],
        )
        return Path(ruta) if ruta else None
    finally:
        root.destroy()


def recolectar_archivos_compatibles(ruta: Union[str, Path]) -> List[Path]:
    """Busca archivos compatibles. Si ruta es carpeta, busca recursivamente."""
    ruta = Path(ruta)
    if ruta.is_file():
        if ruta.suffix.lower() in EXTENSIONES_COMPATIBLES:
            return [ruta]
        print(f"Advertencia: extension no compatible: {ruta}")
        return []

    if not ruta.exists():
        print(f"Ruta no encontrada: {ruta}")
        return []

    archivos = [
        p
        for p in ruta.rglob("*")
        if p.is_file() and p.suffix.lower() in EXTENSIONES_COMPATIBLES
    ]
    return sorted(archivos)


def ruta_salida_unica(ruta: Union[str, Path]) -> Path:
    """Evita sobrescribir archivos cuando SOBRESCRIBIR_SALIDAS=False."""
    ruta = Path(ruta)
    if SOBRESCRIBIR_SALIDAS or not ruta.exists():
        return ruta

    for n in range(2, 10000):
        candidata = ruta.with_name(f"{ruta.stem}_{n}{ruta.suffix}")
        if not candidata.exists():
            return candidata

    raise RuntimeError(f"No se pudo crear un nombre alternativo para: {ruta}")


def nombre_seguro(nombre: Union[str, Path], reemplazo: str = "_") -> str:
    """
    Limpia nombres para carpetas/archivos de salida en Windows.

    No modifica el archivo original. Solo evita problemas como espacios finales,
    puntos finales, caracteres invalidos o nombres reservados.
    """
    nombre = Path(nombre).stem if isinstance(nombre, Path) else str(nombre)
    nombre = re.sub(r'[<>:"/\\|?*\x00-\x1f]', reemplazo, nombre)
    nombre = re.sub(r"\s+", " ", nombre).strip().rstrip(" .")
    nombre = nombre or "archivo"

    if nombre.upper() in NOMBRES_RESERVADOS_WINDOWS:
        nombre = f"{nombre}_archivo"

    return nombre


def carpeta_salida_para(archivo_fuente: Union[str, Path]) -> Path:
    """
    Crea/retorna una subcarpeta por archivo.

    Ejemplo:
    audio.mp3 -> transcripciones/audio/

    Ahi quedan juntos los TXT, SRT, DOCX de transcripcion y la minuta/resumen.
    """
    archivo_fuente = Path(archivo_fuente)
    carpeta = archivo_fuente.parent / CARPETA_SALIDA / nombre_seguro(archivo_fuente.stem)
    carpeta.mkdir(parents=True, exist_ok=True)
    return carpeta


def base_minuta_desde_txt(path_txt: Union[str, Path]) -> str:
    """Evita nombres como archivo_texto_limpio_MINUTA cuando se usa modo 3."""
    stem = Path(path_txt).stem
    for sufijo in ("_texto_limpio", "_con_tiempos", "_MINUTA"):
        if stem.endswith(sufijo):
            stem = stem[: -len(sufijo)]
    return nombre_seguro(stem)


def carpeta_salida_para_txt(path_txt: Union[str, Path]) -> Path:
    """
    Carpeta de salida para el modo 3.

    Si el TXT ya esta en transcripciones/audio/, guarda la minuta ahi.
    Si el TXT esta suelto, crea transcripciones/audio/ junto al TXT.
    """
    path_txt = Path(path_txt)
    base = base_minuta_desde_txt(path_txt)

    if path_txt.parent.name == base and path_txt.parent.parent.name == CARPETA_SALIDA:
        carpeta = path_txt.parent
    elif path_txt.parent.name == CARPETA_SALIDA:
        carpeta = path_txt.parent / base
    else:
        carpeta = path_txt.parent / CARPETA_SALIDA / base

    carpeta.mkdir(parents=True, exist_ok=True)
    return carpeta


def imprimir_salidas(salidas: Dict[str, Union[str, Path]]) -> None:
    print("\nSalidas generadas:")
    for etiqueta, ruta in salidas.items():
        if ruta:
            print(f"- {etiqueta}: {ruta}")


# =============================================================================
# TIEMPOS Y TRANSCRIPCION
# =============================================================================

def formato_hhmmss(segundos: float) -> str:
    segundos = max(0, int(segundos))
    horas = segundos // 3600
    minutos = (segundos % 3600) // 60
    segs = segundos % 60
    return f"{horas:02d}:{minutos:02d}:{segs:02d}"


def formato_srt(segundos: float) -> str:
    total_ms = max(0, int(round(float(segundos) * 1000)))
    horas, resto = divmod(total_ms, 3_600_000)
    minutos, resto = divmod(resto, 60_000)
    segs, ms = divmod(resto, 1000)
    return f"{horas:02d}:{minutos:02d}:{segs:02d},{ms:03d}"


def cargar_modelo_whisper():
    """
    Carga faster-whisper intentando GPU primero.

    La primera vez que uses un modelo puede descargarse. Despues queda en cache
    local y puede funcionar offline.
    """
    global _MODELO_WHISPER_RETENIDO, _MODELO_WHISPER_RETENIDO_INFO, _ULTIMO_WHISPER_INFO

    if _MODELO_WHISPER_RETENIDO is not None:
        info = _MODELO_WHISPER_RETENIDO_INFO or {}
        modelo_retenido = info.get("model_name", "modelo retenido")
        dispositivo = info.get("device", GPU_DEVICE)
        compute_type = info.get("compute_type", GPU_COMPUTE_TYPE)

        if modelo_retenido != WHISPER_MODEL_NAME:
            print(
                "\nAdvertencia: ya hay un modelo Whisper retenido en memoria "
                f"({modelo_retenido}). Se reutilizara para no cargar otro modelo "
                "y evitar presion extra sobre la GPU."
            )
        else:
            print(f"\nReutilizando Whisper retenido en memoria ({modelo_retenido}).")

        return _MODELO_WHISPER_RETENIDO, dispositivo, compute_type

    from faster_whisper import WhisperModel

    print(f"\nCargando Whisper '{WHISPER_MODEL_NAME}' en GPU ({GPU_DEVICE}/{GPU_COMPUTE_TYPE})...")

    try:
        model = WhisperModel(
            WHISPER_MODEL_NAME,
            device=GPU_DEVICE,
            compute_type=GPU_COMPUTE_TYPE,
        )
        print("Whisper cargado en GPU CUDA.")
        _ULTIMO_WHISPER_INFO = {
            "model_name": WHISPER_MODEL_NAME,
            "device": GPU_DEVICE,
            "compute_type": GPU_COMPUTE_TYPE,
        }
        return model, GPU_DEVICE, GPU_COMPUTE_TYPE

    except Exception as exc:
        print("\nADVERTENCIA: no se pudo cargar Whisper en CUDA/GPU.")
        print("Se usara CPU como respaldo. Sera mas lento, pero deberia continuar.")
        print(f"Detalle tecnico: {exc}")

    print(f"\nCargando Whisper en CPU ({CPU_COMPUTE_TYPE}, {CPU_THREADS} hilos)...")
    model = WhisperModel(
        WHISPER_MODEL_NAME,
        device=CPU_DEVICE,
        compute_type=CPU_COMPUTE_TYPE,
        cpu_threads=CPU_THREADS,
    )
    print("Whisper cargado en CPU.")
    _ULTIMO_WHISPER_INFO = {
        "model_name": WHISPER_MODEL_NAME,
        "device": CPU_DEVICE,
        "compute_type": CPU_COMPUTE_TYPE,
    }
    return model, CPU_DEVICE, CPU_COMPUTE_TYPE


def liberar_modelo(model):
    """
    Gestiona la memoria despues de transcribir.

    Por defecto NO fuerza limpieza de CUDA porque en algunos entornos de
    Jupyter/Windows eso puede cerrar el kernel. Si necesitas liberar VRAM antes
    de Ollama, cambia LIBERAR_WHISPER_ANTES_OLLAMA = True.
    """
    global _MODELO_WHISPER_RETENIDO, _MODELO_WHISPER_RETENIDO_INFO, _ULTIMO_WHISPER_INFO

    if not LIBERAR_WHISPER_ANTES_OLLAMA:
        print("\nNo se forzo liberacion de Whisper/CUDA para proteger el kernel de Jupyter.")
        print("Si Ollama queda lento o sin VRAM, prueba WHISPER_MODEL_NAME = \"medium\".")
        if RETENER_MODELO_WHISPER_EN_MEMORIA and model is not None:
            _MODELO_WHISPER_RETENIDO = model
            _MODELO_WHISPER_RETENIDO_INFO = _ULTIMO_WHISPER_INFO
            print("Whisper queda retenido en memoria y se reutilizara si vuelves a transcribir.")
        return model

    print("\nSoltando referencia del modelo Whisper antes de usar Ollama...")
    _MODELO_WHISPER_RETENIDO = None
    _MODELO_WHISPER_RETENIDO_INFO = None
    model = None

    if FORZAR_GC_AL_LIBERAR:
        try:
            gc.collect()
        except Exception:
            pass

    if LIMPIEZA_CUDA_AGRESIVA:
        try:
            import torch

            if torch.cuda.is_available():
                torch.cuda.empty_cache()
                torch.cuda.ipc_collect()
        except Exception as exc:
            print(f"No se pudo ejecutar limpieza CUDA agresiva: {exc}")

    print("Referencia de Whisper soltada. No se ejecuto limpieza CUDA agresiva.")
    return None


def transcribir_archivo(model, archivo: Union[str, Path]) -> Tuple[List[Dict[str, Union[float, str]]], str, object]:
    """Transcribe un archivo y devuelve segmentos, texto limpio e info."""
    archivo = Path(archivo)
    print(f"\nTranscribiendo: {archivo.name}")

    segmentos_generador, info = model.transcribe(
        str(archivo),
        language=LANGUAGE,
        task=TASK,
        beam_size=BEAM_SIZE,
        vad_filter=VAD_FILTER,
    )

    segmentos: List[Dict[str, Union[float, str]]] = []
    for segmento in tqdm(segmentos_generador, desc=f"Segmentos {archivo.name}", unit="seg"):
        texto = (segmento.text or "").strip()
        if not texto:
            continue
        segmentos.append(
            {
                "start": float(segmento.start),
                "end": float(segmento.end),
                "text": texto,
            }
        )

    texto_limpio = " ".join(str(s["text"]).strip() for s in segmentos).strip()
    return segmentos, texto_limpio, info


# =============================================================================
# GUARDADO DE TRANSCRIPCIONES
# =============================================================================

def guardar_txt_con_tiempos(segmentos: Sequence[Dict[str, Union[float, str]]], ruta: Union[str, Path]) -> Path:
    ruta = ruta_salida_unica(ruta)
    ruta.parent.mkdir(parents=True, exist_ok=True)
    lineas = []
    for s in segmentos:
        inicio = formato_hhmmss(float(s["start"]))
        fin = formato_hhmmss(float(s["end"]))
        texto = str(s["text"]).strip()
        lineas.append(f"[{inicio} - {fin}] {texto}")
    ruta.write_text("\n".join(lineas), encoding="utf-8")
    return ruta


def guardar_txt_limpio(texto_limpio: str, ruta: Union[str, Path]) -> Path:
    ruta = ruta_salida_unica(ruta)
    ruta.parent.mkdir(parents=True, exist_ok=True)
    ruta.write_text(texto_limpio.strip() + "\n", encoding="utf-8")
    return ruta


def guardar_srt(segmentos: Sequence[Dict[str, Union[float, str]]], ruta: Union[str, Path]) -> Path:
    ruta = ruta_salida_unica(ruta)
    ruta.parent.mkdir(parents=True, exist_ok=True)
    bloques = []
    for i, s in enumerate(segmentos, start=1):
        inicio = formato_srt(float(s["start"]))
        fin = formato_srt(float(s["end"]))
        texto = str(s["text"]).strip()
        bloques.append(f"{i}\n{inicio} --> {fin}\n{texto}")
    ruta.write_text("\n\n".join(bloques) + "\n", encoding="utf-8")
    return ruta


def _agregar_filas_metadata(doc: Document, pares: Sequence[Tuple[str, str]]) -> None:
    tabla = doc.add_table(rows=0, cols=2)
    tabla.style = "Table Grid"
    for clave, valor in pares:
        fila = tabla.add_row().cells
        fila[0].text = clave
        fila[1].text = valor


def _agregar_texto_largo_docx(doc: Document, texto: str, max_chars: int = 2500) -> None:
    texto = texto.strip()
    if not texto:
        doc.add_paragraph("No se genero texto.")
        return

    for i in range(0, len(texto), max_chars):
        doc.add_paragraph(texto[i : i + max_chars])


def guardar_docx_transcripcion(
    archivo_fuente: Union[str, Path],
    segmentos: Sequence[Dict[str, Union[float, str]]],
    texto_limpio: str,
    ruta: Union[str, Path],
    modelo_whisper: str,
    dispositivo_usado: str,
    compute_type_usado: str,
) -> Path:
    ruta = ruta_salida_unica(ruta)
    ruta.parent.mkdir(parents=True, exist_ok=True)
    archivo_fuente = Path(archivo_fuente)

    doc = Document()
    doc.add_heading("Transcripción automática", level=0)

    _agregar_filas_metadata(
        doc,
        [
            ("Archivo fuente", archivo_fuente.name),
            ("Ruta fuente", str(archivo_fuente)),
            ("Idioma configurado", LANGUAGE),
            ("Modelo Whisper usado", modelo_whisper),
            ("Dispositivo usado", dispositivo_usado),
            ("Tipo de cálculo", compute_type_usado),
        ],
    )

    doc.add_heading("Transcripción con marcas de tiempo", level=1)
    for s in segmentos:
        inicio = formato_hhmmss(float(s["start"]))
        fin = formato_hhmmss(float(s["end"]))
        texto = str(s["text"]).strip()
        doc.add_paragraph(f"[{inicio} - {fin}] {texto}")

    doc.add_heading("Texto limpio", level=1)
    _agregar_texto_largo_docx(doc, texto_limpio)

    doc.save(ruta)
    return ruta


# =============================================================================
# OLLAMA LOCAL
# =============================================================================

def ollama_disponible(modelo: str = OLLAMA_MODEL) -> Tuple[bool, str]:
    """Verifica que Ollama local este activo y que el modelo exista."""
    url = "http://localhost:11434/api/tags"
    try:
        resp = requests.get(url, timeout=10)
        resp.raise_for_status()
        data = resp.json()
    except Exception as exc:
        return (
            False,
            "Ollama no esta disponible en http://localhost:11434. "
            f"Detalle: {exc}",
        )

    modelos = [m.get("name", "") for m in data.get("models", [])]
    if modelo not in modelos:
        return (
            False,
            f"Ollama esta activo, pero no encuentro el modelo '{modelo}'. "
            f"Ejecuta en una terminal: ollama pull {modelo}",
        )

    return True, f"Ollama disponible con modelo '{modelo}'."


def limpiar_respuesta_ollama(texto: str) -> str:
    """Limpia etiquetas de razonamiento o fences si el modelo las agrega."""
    texto = texto or ""
    texto = re.sub(r"<think>.*?</think>", "", texto, flags=re.IGNORECASE | re.DOTALL)
    texto = re.sub(r"^\s*```(?:markdown|md)?\s*", "", texto.strip(), flags=re.IGNORECASE)
    texto = re.sub(r"\s*```\s*$", "", texto.strip())
    return texto.strip()


def llamar_a_ollama(
    prompt: str,
    modelo: str = OLLAMA_MODEL,
    options: Optional[Dict[str, Union[int, float, str]]] = None,
    timeout: int = OLLAMA_TIMEOUT,
) -> str:
    """Llama a Ollama local usando /api/generate sin streaming."""
    url = "http://localhost:11434/api/generate"
    payload = {
        "model": modelo,
        "prompt": prompt,
        "stream": False,
        "options": options or OLLAMA_OPTIONS,
    }

    resp = requests.post(url, json=payload, timeout=timeout)
    if resp.status_code != 200:
        raise RuntimeError(f"Ollama respondio HTTP {resp.status_code}: {resp.text[:1000]}")

    data = resp.json()
    if "error" in data:
        raise RuntimeError(f"Ollama devolvio error: {data['error']}")

    return limpiar_respuesta_ollama(data.get("response", ""))


# =============================================================================
# PROMPTS Y MINUTA
# =============================================================================

REGLAS_PROMPT = """
Reglas estrictas:
- El contenido puede ser una reunion, entrevista, clase, programa de radio,
  cancion, conferencia u otro tipo de audio/video.
- Si NO es una reunion, haz un resumen estructurado del contenido y escribe
  "No especificado" en acuerdos, tareas, responsables, fechas y proximos pasos
  cuando no existan explicitamente.
- Si parece una cancion, resume el tema general o la intencion si es claro,
  pero no reproduzcas la letra completa.
- No inventes participantes.
- No inventes acuerdos.
- No inventes fechas.
- No inventes responsables.
- Si algo no esta claro, escribe "No especificado".
- Distingue entre temas hablados y acuerdos reales.
- No conviertas comentarios generales en acuerdos.
- Mantén tono formal y profesional.
- Corrige errores evidentes de transcripcion solo cuando el contexto sea claro.
- No muestres razonamiento interno.
- No uses etiquetas <think>.
- No agregues informacion externa.
- No hagas suposiciones no sustentadas en la transcripcion.
"""


def dividir_texto_largo(texto: str, tamano: int = TAMANO_FRAGMENTO_MINUTA) -> List[str]:
    """Divide por palabras para no cortar demasiado brusco."""
    texto = re.sub(r"\s+", " ", texto or "").strip()
    if not texto:
        return []
    if len(texto) <= tamano:
        return [texto]

    fragmentos = []
    palabras = texto.split(" ")
    actual: List[str] = []
    longitud_actual = 0

    for palabra in palabras:
        extra = len(palabra) + (1 if actual else 0)
        if actual and longitud_actual + extra > tamano:
            fragmentos.append(" ".join(actual))
            actual = [palabra]
            longitud_actual = len(palabra)
        else:
            actual.append(palabra)
            longitud_actual += extra

    if actual:
        fragmentos.append(" ".join(actual))

    return fragmentos


def crear_prompt_resumen_fragmento(
    fragmento: str,
    numero_fragmento: int,
    total_fragmentos: int,
    nombre_fuente: str,
) -> str:
    return textwrap.dedent(
        f"""
        Eres un asistente local que prepara notas fieles para un resumen/minuta.
        Resume SOLO el contenido del fragmento indicado. No inventes nada.

        Archivo fuente: {nombre_fuente}
        Fragmento: {numero_fragmento} de {total_fragmentos}

        {REGLAS_PROMPT}

        Devuelve notas en Markdown con estas secciones:
        ## Temas mencionados
        - ...

        ## Posibles acuerdos explicitos
        - ...

        ## Tareas explicitas
        - Responsable: No especificado | Tarea: ... | Fecha limite: No especificado

        ## Pendientes o dudas
        - ...

        Transcripcion del fragmento:
        \"\"\"
        {fragmento}
        \"\"\"
        """
    ).strip()


def crear_prompt_minuta_final(
    material: Union[str, Sequence[str]],
    nombre_fuente: str,
) -> str:
    if isinstance(material, str):
        material_texto = material
    else:
        material_texto = "\n\n".join(material)

    return textwrap.dedent(
        f"""
        Eres un asistente local que redacta un resumen/minuta formal a partir de
        una transcripcion o de resumenes intermedios. Usa unicamente el material
        dado. No asumas que el contenido es una reunion.

        Archivo fuente: {nombre_fuente}

        {REGLAS_PROMPT}

        Devuelve EXACTAMENTE esta estructura en Markdown, sin secciones extra:

        # Resumen / minuta automática

        ## Archivo fuente
        {nombre_fuente}

        ## Resumen ejecutivo
        Resumen breve del contenido. Si es claro, indica si parece reunion,
        entrevista, clase, programa de radio, cancion u otro formato.

        ## Temas tratados
        - ...

        ## Acuerdos
        - ...

        ## Tareas y responsables
        | Responsable | Tarea | Fecha límite |
        |---|---|---|
        | No especificado | ... | No especificado |

        ## Puntos pendientes
        - ...

        ## Riesgos, dudas o temas por aclarar
        - ...

        ## Próximos pasos sugeridos
        - ...

        ## Nota de cautela
        Este resumen/minuta fue generado automáticamente a partir de una transcripción. Debe ser revisado antes de usarse como documento oficial.

        Material disponible:
        \"\"\"
        {material_texto}
        \"\"\"
        """
    ).strip()


def resumir_fragmentos(texto: str, nombre_fuente: str) -> List[str]:
    fragmentos = dividir_texto_largo(texto, TAMANO_FRAGMENTO_MINUTA)
    if not fragmentos:
        return []

    resumenes = []
    for i, fragmento in enumerate(
        tqdm(fragmentos, desc="Resumiendo fragmentos", unit="frag"),
        start=1,
    ):
        prompt = crear_prompt_resumen_fragmento(fragmento, i, len(fragmentos), nombre_fuente)
        respuesta = llamar_a_ollama(prompt)
        resumenes.append(f"## Fragmento {i}\n{respuesta}")

    return resumenes


def consolidar_minuta_final(material: Union[str, Sequence[str]], nombre_fuente: str) -> str:
    prompt = crear_prompt_minuta_final(material, nombre_fuente)
    return llamar_a_ollama(prompt)


def generar_minuta_de_texto(texto_limpio: str, nombre_fuente: str) -> str:
    """Genera minuta; si el texto es largo, resume fragmentos y consolida."""
    fragmentos = dividir_texto_largo(texto_limpio, TAMANO_FRAGMENTO_MINUTA)
    if not fragmentos:
        raise ValueError("El texto esta vacio; no se puede generar minuta.")

    if len(fragmentos) == 1:
        return consolidar_minuta_final(fragmentos[0], nombre_fuente)

    resumenes = resumir_fragmentos(texto_limpio, nombre_fuente)
    return consolidar_minuta_final(resumenes, nombre_fuente)


# =============================================================================
# GUARDADO DE MINUTA
# =============================================================================

def guardar_minuta_txt(minuta_markdown: str, carpeta: Union[str, Path], base_nombre: str) -> Path:
    carpeta = Path(carpeta)
    carpeta.mkdir(parents=True, exist_ok=True)
    base_nombre = nombre_seguro(base_nombre)
    ruta = ruta_salida_unica(carpeta / f"{base_nombre}_MINUTA.txt")
    ruta.write_text(minuta_markdown.strip() + "\n", encoding="utf-8")
    return ruta


def _es_linea_tabla(linea: str) -> bool:
    linea = linea.strip()
    return linea.startswith("|") and linea.endswith("|") and linea.count("|") >= 2


def _parsear_fila_tabla(linea: str) -> List[str]:
    return [celda.strip() for celda in linea.strip().strip("|").split("|")]


def _es_separador_markdown(linea: str) -> bool:
    celdas = _parsear_fila_tabla(linea)
    if not celdas:
        return False
    return all(re.fullmatch(r":?-{3,}:?", celda.strip()) for celda in celdas)


def _agregar_tabla_markdown(doc: Document, lineas_tabla: Sequence[str]) -> bool:
    filas = [_parsear_fila_tabla(l) for l in lineas_tabla if _es_linea_tabla(l)]
    if len(filas) < 2:
        return False

    if len(filas) >= 2 and _es_separador_markdown(lineas_tabla[1]):
        filas_utiles = [filas[0]] + filas[2:]
    else:
        filas_utiles = filas

    if not filas_utiles:
        return False

    num_cols = max(len(f) for f in filas_utiles)
    tabla = doc.add_table(rows=0, cols=num_cols)
    tabla.style = "Table Grid"
    tabla.alignment = WD_TABLE_ALIGNMENT.CENTER

    for fila_datos in filas_utiles:
        celdas = tabla.add_row().cells
        for i in range(num_cols):
            celdas[i].text = fila_datos[i] if i < len(fila_datos) else ""

    return True


def _agregar_markdown_basico_a_docx(doc: Document, markdown: str) -> None:
    lineas = markdown.splitlines()
    i = 0

    while i < len(lineas):
        linea = lineas[i].rstrip()
        limpia = linea.strip()

        if not limpia:
            i += 1
            continue

        if _es_linea_tabla(limpia):
            bloque = []
            while i < len(lineas) and _es_linea_tabla(lineas[i].strip()):
                bloque.append(lineas[i].strip())
                i += 1

            if not _agregar_tabla_markdown(doc, bloque):
                for l in bloque:
                    doc.add_paragraph(l)
            continue

        if limpia.startswith("# "):
            doc.add_heading(limpia[2:].strip(), level=0)
        elif limpia.startswith("## "):
            doc.add_heading(limpia[3:].strip(), level=1)
        elif limpia.startswith("### "):
            doc.add_heading(limpia[4:].strip(), level=2)
        elif limpia.startswith("- "):
            doc.add_paragraph(limpia[2:].strip(), style="List Bullet")
        else:
            doc.add_paragraph(limpia)

        i += 1


def guardar_minuta_docx(minuta_markdown: str, carpeta: Union[str, Path], base_nombre: str) -> Path:
    carpeta = Path(carpeta)
    carpeta.mkdir(parents=True, exist_ok=True)
    base_nombre = nombre_seguro(base_nombre)
    ruta = ruta_salida_unica(carpeta / f"{base_nombre}_MINUTA.docx")

    doc = Document()
    _agregar_markdown_basico_a_docx(doc, minuta_markdown)
    doc.save(ruta)
    return ruta


# =============================================================================
# FLUJOS PRINCIPALES
# =============================================================================

def generar_minuta_desde_txt() -> Dict[str, object]:
    """Modo 3: genera minuta desde un TXT limpio ya existente."""
    resultado: Dict[str, object] = {
        "modo": "minuta_desde_txt",
        "archivo": None,
        "minuta_generada": False,
        "ollama_disponible": False,
        "salidas": {},
        "error": None,
    }

    archivo_txt = seleccionar_txt_limpio()
    if not archivo_txt:
        resultado["error"] = "No se selecciono ningun archivo TXT."
        print(resultado["error"])
        return resultado

    resultado["archivo"] = str(archivo_txt)
    ok, mensaje = ollama_disponible()
    print(mensaje)
    resultado["ollama_disponible"] = ok
    if not ok:
        resultado["error"] = mensaje
        return resultado

    try:
        texto = archivo_txt.read_text(encoding="utf-8", errors="replace").strip()
        base = base_minuta_desde_txt(archivo_txt)
        minuta = generar_minuta_de_texto(texto, archivo_txt.name)

        carpeta = carpeta_salida_para_txt(archivo_txt)
        ruta_txt = guardar_minuta_txt(minuta, carpeta, base)
        ruta_docx = guardar_minuta_docx(minuta, carpeta, base)

        resultado["minuta_generada"] = True
        resultado["salidas"] = {
            "minuta_txt": str(ruta_txt),
            "minuta_docx": str(ruta_docx),
        }
        imprimir_salidas(resultado["salidas"])  # type: ignore[arg-type]
        return resultado

    except Exception as exc:
        resultado["error"] = str(exc)
        print(f"Error generando minuta desde TXT: {exc}")
        return resultado


def procesar_un_archivo(
    archivo: Union[str, Path],
    model,
    dispositivo_usado: str,
    compute_type_usado: str,
    generar_minuta: bool = False,
    ollama_ok: bool = False,
) -> Dict[str, object]:
    """Transcribe un archivo y, opcionalmente, genera minuta."""
    archivo = Path(archivo)
    salida_dir = carpeta_salida_para(archivo)
    base = nombre_seguro(archivo.stem)

    resultado: Dict[str, object] = {
        "archivo": str(archivo),
        "modelo_whisper": WHISPER_MODEL_NAME,
        "dispositivo": dispositivo_usado,
        "compute_type": compute_type_usado,
        "minuta_generada": False,
        "ollama_disponible": ollama_ok,
        "salidas": {},
        "error": None,
    }

    try:
        segmentos, texto_limpio, _info = transcribir_archivo(model, archivo)

        ruta_tiempos = guardar_txt_con_tiempos(
            segmentos,
            salida_dir / f"{base}_con_tiempos.txt",
        )
        ruta_limpio = guardar_txt_limpio(
            texto_limpio,
            salida_dir / f"{base}_texto_limpio.txt",
        )
        ruta_srt = guardar_srt(segmentos, salida_dir / f"{base}.srt")
        ruta_docx = guardar_docx_transcripcion(
            archivo_fuente=archivo,
            segmentos=segmentos,
            texto_limpio=texto_limpio,
            ruta=salida_dir / f"{base}.docx",
            modelo_whisper=WHISPER_MODEL_NAME,
            dispositivo_usado=dispositivo_usado,
            compute_type_usado=compute_type_usado,
        )

        salidas: Dict[str, str] = {
            "txt_con_tiempos": str(ruta_tiempos),
            "txt_limpio": str(ruta_limpio),
            "srt": str(ruta_srt),
            "docx_transcripcion": str(ruta_docx),
        }

        if generar_minuta and ollama_ok:
            minuta = generar_minuta_de_texto(texto_limpio, archivo.name)
            ruta_minuta_txt = guardar_minuta_txt(minuta, salida_dir, base)
            ruta_minuta_docx = guardar_minuta_docx(minuta, salida_dir, base)
            salidas["minuta_txt"] = str(ruta_minuta_txt)
            salidas["minuta_docx"] = str(ruta_minuta_docx)
            resultado["minuta_generada"] = True

        resultado["salidas"] = salidas
        imprimir_salidas(salidas)
        return resultado

    except Exception as exc:
        resultado["error"] = str(exc)
        print(f"\nERROR procesando {archivo.name}: {exc}")
        print("Se continuara con los demas archivos.")
        return resultado


def _generar_minuta_para_resultado(resultado: Dict[str, object]) -> Dict[str, object]:
    """Genera minuta despues de liberar Whisper, usando el TXT limpio guardado."""
    if resultado.get("error"):
        return resultado

    salidas = dict(resultado.get("salidas") or {})
    txt_limpio = salidas.get("txt_limpio")
    archivo_fuente = Path(str(resultado.get("archivo", "")))

    if not txt_limpio:
        resultado["error"] = "No hay TXT limpio para generar minuta."
        return resultado

    try:
        txt_path = Path(str(txt_limpio))
        texto = txt_path.read_text(encoding="utf-8", errors="replace").strip()
        base = nombre_seguro(archivo_fuente.stem) if archivo_fuente.name else base_minuta_desde_txt(txt_path)
        minuta = generar_minuta_de_texto(texto, archivo_fuente.name or txt_path.name)

        carpeta = txt_path.parent
        ruta_minuta_txt = guardar_minuta_txt(minuta, carpeta, base)
        ruta_minuta_docx = guardar_minuta_docx(minuta, carpeta, base)

        salidas["minuta_txt"] = str(ruta_minuta_txt)
        salidas["minuta_docx"] = str(ruta_minuta_docx)
        resultado["salidas"] = salidas
        resultado["minuta_generada"] = True
        resultado["ollama_disponible"] = True
        imprimir_salidas(
            {
                "minuta_txt": ruta_minuta_txt,
                "minuta_docx": ruta_minuta_docx,
            }
        )
        return resultado

    except Exception as exc:
        resultado["minuta_generada"] = False
        resultado["error_minuta"] = str(exc)
        print(f"No se pudo generar minuta para {archivo_fuente.name}: {exc}")
        return resultado


def transcribir(
    archivos: Optional[Sequence[Union[str, Path]]] = None,
    generar_minuta: bool = GENERAR_MINUTA,
) -> List[Dict[str, object]]:
    """
    Transcribe una lista de archivos.

    Si generar_minuta=True, primero transcribe todo, libera Whisper y luego llama
    a Ollama para reducir consumo simultaneo de VRAM.
    """
    if archivos is None:
        tipo, ruta = elegir_archivo_o_carpeta()
        if not ruta:
            print("No se selecciono ruta.")
            return []
        archivos = recolectar_archivos_compatibles(ruta) if tipo == "carpeta" else recolectar_archivos_compatibles(ruta)

    archivos = [Path(a) for a in archivos]
    if not archivos:
        print("No se encontraron archivos compatibles.")
        return []

    resultados: List[Dict[str, object]] = []

    try:
        model, dispositivo_usado, compute_type_usado = cargar_modelo_whisper()
    except Exception as exc:
        print(f"No se pudo cargar Whisper ni en GPU ni en CPU: {exc}")
        return [
            {
                "archivo": str(a),
                "modelo_whisper": WHISPER_MODEL_NAME,
                "dispositivo": None,
                "compute_type": None,
                "minuta_generada": False,
                "ollama_disponible": False,
                "salidas": {},
                "error": str(exc),
            }
            for a in archivos
        ]

    try:
        for archivo in tqdm(archivos, desc="Archivos", unit="archivo"):
            resultado = procesar_un_archivo(
                archivo=archivo,
                model=model,
                dispositivo_usado=dispositivo_usado,
                compute_type_usado=compute_type_usado,
                generar_minuta=False,
                ollama_ok=False,
            )
            resultados.append(resultado)
    finally:
        model = liberar_modelo(model)

    if generar_minuta:
        print("\nVerificando Ollama local para generar minutas...")
        ok, mensaje = ollama_disponible()
        print(mensaje)

        if not ok:
            print("Las transcripciones ya quedaron guardadas. Puedes iniciar Ollama y usar el modo 3.")
            for r in resultados:
                r["ollama_disponible"] = False
            return resultados

        for resultado in tqdm(resultados, desc="Generando minutas", unit="archivo"):
            _generar_minuta_para_resultado(resultado)

    return resultados


def main() -> List[Dict[str, object]]:
    """
    Punto de entrada pensado para Jupyter:

    resultados = main()
    """
    print("Herramienta local de transcripcion y minutas")
    print("Todo corre localmente: faster-whisper y Ollama en localhost.")
    print("Sugerencia RTX 3070 8 GB: para archivos largos puedes usar modo 1 y luego modo 3.")

    modo = elegir_modo_proceso()

    if modo == 3:
        return [generar_minuta_desde_txt()]

    tipo, ruta = elegir_archivo_o_carpeta()
    if not ruta:
        print("No se selecciono ruta.")
        return []

    archivos = recolectar_archivos_compatibles(ruta)
    if not archivos:
        print("No se encontraron archivos compatibles.")
        return []

    generar_minuta = modo == 2
    return transcribir(archivos=archivos, generar_minuta=generar_minuta)


# En Jupyter ejecuta en una celda aparte:
# resultados = main()
