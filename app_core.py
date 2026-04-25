# -*- coding: utf-8 -*-
"""
Nucleo de procesamiento para Gamboa Transcriptor.

Este modulo no depende de PySide6. La interfaz grafica llama estas funciones
desde un hilo de trabajo para mantener la ventana activa mientras se transcribe.
"""

from __future__ import annotations

import gc
import json
import os
import re
import shutil
import subprocess
import time
import traceback
from dataclasses import dataclass, field
from pathlib import Path
from typing import Callable, Dict, List, Optional, Sequence, Tuple, Union

import requests
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT


# =============================================================================
# CONSTANTES
# =============================================================================

APP_NAME = "Gamboa Desarrollos - Transcriptor Local y Minutador"
APP_SHORT_NAME = "Gamboa Transcriptor"

EXTENSIONES_AUDIO = {
    ".mp3",
    ".wav",
    ".m4a",
    ".aac",
    ".flac",
    ".ogg",
    ".wma",
}

EXTENSIONES_VIDEO = {
    ".mp4",
    ".mkv",
    ".mov",
    ".avi",
    ".webm",
    ".m4v",
    ".wmv",
}

EXTENSIONES_COMPATIBLES = EXTENSIONES_AUDIO | EXTENSIONES_VIDEO

IDIOMAS_MINUTA = {
    "source": ("Idioma de la transcripcion", ""),
    "es": ("Español", "es"),
    "en": ("Ingles", "en"),
    "zh": ("Chino", "zh"),
    "pt": ("Portugues", "pt"),
    "fr": ("Frances", "fr"),
    "de": ("Aleman", "de"),
    "it": ("Italiano", "it"),
}

TODOS_IDIOMAS_MINUTA = ["es", "en", "zh", "pt"]

NOMBRES_RESERVADOS_WINDOWS = {
    "CON",
    "PRN",
    "AUX",
    "NUL",
    *(f"COM{i}" for i in range(1, 10)),
    *(f"LPT{i}" for i in range(1, 10)),
}


# =============================================================================
# CONFIGURACION Y RESULTADOS
# =============================================================================

@dataclass
class OutputOptions:
    txt_con_tiempos: bool = True
    txt_limpio: bool = True
    srt: bool = True
    docx_transcripcion: bool = True
    docx_minuta: bool = True


@dataclass
class AppConfig:
    whisper_model_name: str = "large-v3"
    # Use None para que faster-whisper detecte el idioma automaticamente.
    language: Optional[str] = "es"
    task: str = "transcribe"
    beam_size: int = 5
    vad_filter: bool = True

    device_mode: str = "auto"  # auto, cuda, cpu
    gpu_device: str = "cuda"
    gpu_compute_type: str = "float16"
    cpu_device: str = "cpu"
    cpu_compute_type: str = "int8"
    cpu_threads: int = 6

    ollama_model: str = "qwen3:8b"
    ollama_url: str = "http://localhost:11434/api/generate"
    ollama_tags_url: str = "http://localhost:11434/api/tags"
    ollama_timeout: int = 900
    ollama_options: Dict[str, Union[int, float, str]] = field(
        default_factory=lambda: {"temperature": 0.1, "num_ctx": 4096}
    )
    tamano_fragmento_minuta: int = 8000
    idioma_salida_minuta: str = "es"  # source, es, en, zh, pt, fr, de, it, all

    output_root: Optional[Path] = None
    usar_carpeta_automatica: bool = True
    sobrescribir_salidas: bool = False

    liberar_whisper_antes_ollama: bool = False
    limpieza_cuda_agresiva: bool = False

    output_options: OutputOptions = field(default_factory=OutputOptions)


@dataclass
class ProcessResult:
    source: str
    ok: bool = False
    error: Optional[str] = None
    output_dir: Optional[str] = None
    outputs: Dict[str, str] = field(default_factory=dict)
    minuta_generada: bool = False
    dispositivo: Optional[str] = None
    compute_type: Optional[str] = None
    idioma_configurado: Optional[str] = None
    idioma_detectado: Optional[str] = None
    probabilidad_idioma: Optional[float] = None
    idiomas_minuta: List[str] = field(default_factory=list)
    elapsed_seconds: float = 0.0


@dataclass
class ProcessSummary:
    results: List[ProcessResult]
    output_root: Optional[str]
    elapsed_seconds: float
    cancelled: bool = False

    @property
    def ok_count(self) -> int:
        return sum(1 for r in self.results if r.ok)

    @property
    def error_count(self) -> int:
        return sum(1 for r in self.results if not r.ok)

    @property
    def minute_count(self) -> int:
        return sum(1 for r in self.results if r.minuta_generada)


@dataclass
class ProcessCallbacks:
    log: Callable[[str], None] = lambda _msg: None
    status: Callable[[str], None] = lambda _msg: None
    general_progress: Callable[[int, int], None] = lambda _done, _total: None
    file_progress: Callable[[int], None] = lambda _percent: None
    current_file: Callable[[str], None] = lambda _path: None
    error: Callable[[str], None] = lambda _msg: None
    should_stop_after_current: Callable[[], bool] = lambda: False


# =============================================================================
# UTILIDADES
# =============================================================================

def now_label() -> str:
    return time.strftime("%Y-%m-%d %H:%M:%S")


def log_line(msg: str) -> str:
    return f"[{now_label()}] {msg}"


def nombre_seguro(nombre: Union[str, Path], reemplazo: str = "_") -> str:
    nombre = Path(nombre).stem if isinstance(nombre, Path) else str(nombre)
    nombre = re.sub(r'[<>:"/\\|?*\x00-\x1f]', reemplazo, nombre)
    nombre = re.sub(r"\s+", " ", nombre).strip().rstrip(" .")
    nombre = nombre or "archivo"
    if nombre.upper() in NOMBRES_RESERVADOS_WINDOWS:
        nombre = f"{nombre}_archivo"
    return nombre


def ruta_salida_unica(ruta: Union[str, Path], sobrescribir: bool = False) -> Path:
    ruta = Path(ruta)
    if sobrescribir or not ruta.exists():
        return ruta

    for n in range(2, 10000):
        candidata = ruta.with_name(f"{ruta.stem}_{n}{ruta.suffix}")
        if not candidata.exists():
            return candidata

    raise RuntimeError(f"No se pudo crear nombre alternativo para: {ruta}")


def formato_hhmmss(segundos: float) -> str:
    segundos = max(0, int(segundos))
    horas = segundos // 3600
    minutos = (segundos % 3600) // 60
    segs = segundos % 60
    return f"{horas:02d}:{minutos:02d}:{segs:02d}"


def formato_srt(segundos: float) -> str:
    total_ms = max(0, int(round(float(segundos) * 1000)))
    horas, resto = divmod(total_ms, 3_600_000)
    minutos, resto = divmod(resto, 60_000)
    segs, ms = divmod(resto, 1000)
    return f"{horas:02d}:{minutos:02d}:{segs:02d},{ms:03d}"


def etiqueta_idioma_configurado(language: Optional[str]) -> str:
    if language is None:
        return "Auto detectar"
    etiquetas = {
        "es": "es - Español",
        "en": "en - Ingles",
        "pt": "pt - Portugues",
        "fr": "fr - Frances",
        "de": "de - Aleman",
        "it": "it - Italiano",
    }
    return etiquetas.get(language, language)


def etiqueta_idioma_minuta(codigo: str) -> str:
    return IDIOMAS_MINUTA.get(codigo, (codigo, codigo))[0]


def sufijo_idioma_minuta(codigo: str) -> str:
    if codigo == "source":
        return ""
    return IDIOMAS_MINUTA.get(codigo, (codigo, codigo))[1] or codigo


def resolver_idiomas_minuta(config: AppConfig, idioma_detectado: Optional[str] = None) -> List[str]:
    seleccionado = (config.idioma_salida_minuta or "es").lower()
    if seleccionado == "all":
        return TODOS_IDIOMAS_MINUTA.copy()
    if seleccionado == "source":
        detectado = (idioma_detectado or "").lower().strip()
        if detectado:
            return [detectado]
        return ["source"]
    return [seleccionado]


def recolectar_archivos_compatibles(ruta: Union[str, Path]) -> List[Path]:
    ruta = Path(ruta)
    if ruta.is_file():
        return [ruta] if ruta.suffix.lower() in EXTENSIONES_COMPATIBLES else []
    if not ruta.exists():
        return []
    return sorted(
        p
        for p in ruta.rglob("*")
        if p.is_file() and p.suffix.lower() in EXTENSIONES_COMPATIBLES
    )


def base_desde_txt(path_txt: Union[str, Path]) -> str:
    stem = Path(path_txt).stem
    for sufijo in ("_texto_limpio", "_con_tiempos", "_MINUTA"):
        if stem.endswith(sufijo):
            stem = stem[: -len(sufijo)]
    return nombre_seguro(stem)


def resolver_output_root(source: Path, config: AppConfig) -> Path:
    if config.usar_carpeta_automatica or config.output_root is None:
        return source.parent / "transcripciones"
    return Path(config.output_root)


def output_dir_para_archivo(source: Path, config: AppConfig) -> Path:
    root = resolver_output_root(source, config)
    carpeta = root / nombre_seguro(source.stem)
    carpeta.mkdir(parents=True, exist_ok=True)
    return carpeta


def output_dir_para_txt(path_txt: Path, config: AppConfig) -> Path:
    base = base_desde_txt(path_txt)
    if config.usar_carpeta_automatica or config.output_root is None:
        if path_txt.parent.name == base and path_txt.parent.parent.name == "transcripciones":
            carpeta = path_txt.parent
        else:
            carpeta = path_txt.parent / "transcripciones" / base
    else:
        carpeta = Path(config.output_root) / base
    carpeta.mkdir(parents=True, exist_ok=True)
    return carpeta


def abrir_carpeta(path: Union[str, Path]) -> None:
    path = Path(path)
    if path.exists():
        os.startfile(str(path))


# =============================================================================
# VERIFICACIONES
# =============================================================================

def verificar_ffmpeg() -> Dict[str, object]:
    exe = shutil.which("ffmpeg")
    return {
        "available": exe is not None,
        "path": exe,
        "message": "FFmpeg disponible." if exe else "FFmpeg no detectado. PyAV puede abrir muchos formatos sin FFmpeg externo.",
    }


def verificar_cuda() -> Dict[str, object]:
    result: Dict[str, object] = {
        "available": False,
        "gpu_name": "No detectada",
        "memory_total_mb": None,
        "message": "CUDA no disponible. Se usara CPU.",
    }

    try:
        import ctranslate2

        count = ctranslate2.get_cuda_device_count()
        if count and count > 0:
            result["available"] = True
            result["message"] = "CUDA disponible para CTranslate2/faster-whisper."
    except Exception as exc:
        result["message"] = f"No se pudo verificar CUDA con CTranslate2: {exc}"

    try:
        cmd = [
            "nvidia-smi",
            "--query-gpu=name,memory.total",
            "--format=csv,noheader,nounits",
        ]
        completed = subprocess.run(cmd, capture_output=True, text=True, timeout=5)
        if completed.returncode == 0 and completed.stdout.strip():
            first = completed.stdout.strip().splitlines()[0]
            parts = [p.strip() for p in first.split(",")]
            result["gpu_name"] = parts[0] if parts else "NVIDIA GPU"
            if len(parts) > 1:
                try:
                    result["memory_total_mb"] = int(parts[1])
                except ValueError:
                    pass
            if result["available"]:
                mem = result["memory_total_mb"]
                result["message"] = f"CUDA disponible. GPU: {result['gpu_name']} ({mem} MB VRAM)."
            else:
                result["message"] = f"GPU NVIDIA detectada: {result['gpu_name']}, pero CUDA no fue confirmada por CTranslate2."
    except Exception:
        pass

    return result


def verificar_ollama(modelo: str = "qwen3:8b", tags_url: str = "http://localhost:11434/api/tags") -> Dict[str, object]:
    try:
        resp = requests.get(tags_url, timeout=8)
        resp.raise_for_status()
        data = resp.json()
    except Exception as exc:
        return {
            "available": False,
            "model_installed": False,
            "models": [],
            "message": f"Ollama no esta activo. Puede transcribir, pero no generar minutas. Detalle: {exc}",
        }

    modelos = [m.get("name", "") for m in data.get("models", [])]
    installed = modelo in modelos
    if installed:
        msg = f"Ollama disponible y modelo {modelo} instalado."
    else:
        msg = f"Ollama activo, pero el modelo {modelo} no aparece instalado. Ejecute: ollama pull {modelo}"
    return {
        "available": True,
        "model_installed": installed,
        "models": modelos,
        "message": msg,
    }


def verificar_whisper_import() -> Dict[str, object]:
    try:
        import faster_whisper  # noqa: F401

        return {"available": True, "message": "faster-whisper disponible."}
    except Exception as exc:
        return {"available": False, "message": f"faster-whisper no esta disponible: {exc}"}


# =============================================================================
# EXPORTACION DE TRANSCRIPCION
# =============================================================================

def guardar_txt_con_tiempos(
    segmentos: Sequence[Dict[str, Union[float, str]]],
    ruta: Union[str, Path],
    sobrescribir: bool,
) -> Path:
    ruta = ruta_salida_unica(ruta, sobrescribir)
    ruta.parent.mkdir(parents=True, exist_ok=True)
    lineas = []
    for s in segmentos:
        inicio = formato_hhmmss(float(s["start"]))
        fin = formato_hhmmss(float(s["end"]))
        lineas.append(f"[{inicio} - {fin}] {str(s['text']).strip()}")
    ruta.write_text("\n".join(lineas) + "\n", encoding="utf-8")
    return ruta


def guardar_txt_limpio(texto_limpio: str, ruta: Union[str, Path], sobrescribir: bool) -> Path:
    ruta = ruta_salida_unica(ruta, sobrescribir)
    ruta.parent.mkdir(parents=True, exist_ok=True)
    ruta.write_text(texto_limpio.strip() + "\n", encoding="utf-8")
    return ruta


def guardar_srt(
    segmentos: Sequence[Dict[str, Union[float, str]]],
    ruta: Union[str, Path],
    sobrescribir: bool,
) -> Path:
    ruta = ruta_salida_unica(ruta, sobrescribir)
    ruta.parent.mkdir(parents=True, exist_ok=True)
    bloques = []
    for i, s in enumerate(segmentos, start=1):
        inicio = formato_srt(float(s["start"]))
        fin = formato_srt(float(s["end"]))
        bloques.append(f"{i}\n{inicio} --> {fin}\n{str(s['text']).strip()}")
    ruta.write_text("\n\n".join(bloques) + "\n", encoding="utf-8")
    return ruta


def _agregar_filas_metadata(doc: Document, pares: Sequence[Tuple[str, str]]) -> None:
    tabla = doc.add_table(rows=0, cols=2)
    tabla.style = "Table Grid"
    for clave, valor in pares:
        fila = tabla.add_row().cells
        fila[0].text = clave
        fila[1].text = valor


def _agregar_texto_largo_docx(doc: Document, texto: str, max_chars: int = 2500) -> None:
    texto = texto.strip()
    if not texto:
        doc.add_paragraph("No se genero texto.")
        return
    for i in range(0, len(texto), max_chars):
        doc.add_paragraph(texto[i : i + max_chars])


def guardar_docx_transcripcion(
    archivo_fuente: Union[str, Path],
    segmentos: Sequence[Dict[str, Union[float, str]]],
    texto_limpio: str,
    ruta: Union[str, Path],
    modelo_whisper: str,
    idioma_configurado: str,
    idioma_detectado: str,
    probabilidad_idioma: Optional[float],
    dispositivo_usado: str,
    compute_type_usado: str,
    sobrescribir: bool,
) -> Path:
    ruta = ruta_salida_unica(ruta, sobrescribir)
    ruta.parent.mkdir(parents=True, exist_ok=True)
    archivo_fuente = Path(archivo_fuente)

    doc = Document()
    doc.add_heading("Transcripcion automatica", level=0)
    _agregar_filas_metadata(
        doc,
        [
            ("Archivo fuente", archivo_fuente.name),
            ("Ruta fuente", str(archivo_fuente)),
            ("Idioma configurado", idioma_configurado),
            ("Idioma detectado", idioma_detectado),
            ("Probabilidad idioma", f"{probabilidad_idioma:.2%}" if probabilidad_idioma is not None else "No especificado"),
            ("Modelo Whisper usado", modelo_whisper),
            ("Dispositivo usado", dispositivo_usado),
            ("Tipo de calculo", compute_type_usado),
        ],
    )
    doc.add_heading("Transcripcion con marcas de tiempo", level=1)
    for s in segmentos:
        inicio = formato_hhmmss(float(s["start"]))
        fin = formato_hhmmss(float(s["end"]))
        doc.add_paragraph(f"[{inicio} - {fin}] {str(s['text']).strip()}")
    doc.add_heading("Texto limpio", level=1)
    _agregar_texto_largo_docx(doc, texto_limpio)
    doc.save(ruta)
    return ruta


# =============================================================================
# OLLAMA Y MINUTAS
# =============================================================================

REGLAS_PROMPT = """
Reglas estrictas:
- No inventes participantes.
- No inventes acuerdos.
- No inventes fechas.
- No inventes responsables.
- Si algo no esta claro, escribe "No especificado".
- Distingue entre temas hablados y acuerdos reales.
- No conviertas comentarios generales en acuerdos.
- Mantén tono formal y profesional.
- Corrige errores evidentes de transcripcion solo cuando el contexto sea claro.
- No muestres razonamiento interno.
- No agregues informacion externa.
- No hagas suposiciones no sustentadas en la transcripcion.
- Si el contenido no es una reunion, no inventes acuerdos ni tareas; usa "No especificado" donde aplique.
"""


def limpiar_respuesta_ollama(texto: str) -> str:
    texto = texto or ""
    texto = re.sub(r"<think>.*?</think>", "", texto, flags=re.IGNORECASE | re.DOTALL)
    texto = re.sub(r"^\s*```(?:markdown|md)?\s*", "", texto.strip(), flags=re.IGNORECASE)
    texto = re.sub(r"\s*```\s*$", "", texto.strip())
    return texto.strip()


def llamar_a_ollama(prompt: str, config: AppConfig) -> str:
    payload = {
        "model": config.ollama_model,
        "prompt": prompt,
        "stream": False,
        "options": config.ollama_options,
    }
    resp = requests.post(config.ollama_url, json=payload, timeout=config.ollama_timeout)
    if resp.status_code != 200:
        raise RuntimeError(f"Ollama respondio HTTP {resp.status_code}: {resp.text[:1000]}")
    data = resp.json()
    if "error" in data:
        raise RuntimeError(f"Ollama devolvio error: {data['error']}")
    return limpiar_respuesta_ollama(data.get("response", ""))


def dividir_texto_largo(texto: str, tamano: int) -> List[str]:
    texto = re.sub(r"\s+", " ", texto or "").strip()
    if not texto:
        return []
    if len(texto) <= tamano:
        return [texto]

    fragmentos: List[str] = []
    palabras = texto.split(" ")
    actual: List[str] = []
    longitud = 0
    for palabra in palabras:
        extra = len(palabra) + (1 if actual else 0)
        if actual and longitud + extra > tamano:
            fragmentos.append(" ".join(actual))
            actual = [palabra]
            longitud = len(palabra)
        else:
            actual.append(palabra)
            longitud += extra
    if actual:
        fragmentos.append(" ".join(actual))
    return fragmentos


def crear_prompt_resumen_fragmento(
    fragmento: str,
    i: int,
    total: int,
    nombre_fuente: str,
    idioma_salida: str,
) -> str:
    etiqueta = etiqueta_idioma_minuta(idioma_salida)
    instruccion_idioma = (
        "Redacta la respuesta en el mismo idioma predominante de la transcripcion."
        if idioma_salida == "source"
        else f"Redacta la respuesta final en {etiqueta}."
    )
    return f"""
Eres un asistente local que prepara notas fieles para una minuta/resumen.
Resume SOLO el contenido del fragmento indicado. No inventes nada.

Archivo fuente: {nombre_fuente}
Fragmento: {i} de {total}
Idioma de salida: {etiqueta}
{instruccion_idioma}

{REGLAS_PROMPT}

Devuelve notas en Markdown con estas secciones:
## Temas mencionados
- ...

## Posibles acuerdos explicitos
- ...

## Tareas explicitas
- Responsable: No especificado | Tarea: ... | Fecha limite: No especificado

## Pendientes o dudas
- ...

Transcripcion del fragmento:
\"\"\"
{fragmento}
\"\"\"
""".strip()


def crear_prompt_minuta_final(
    material: Union[str, Sequence[str]],
    nombre_fuente: str,
    idioma_salida: str,
) -> str:
    material_texto = material if isinstance(material, str) else "\n\n".join(material)
    etiqueta = etiqueta_idioma_minuta(idioma_salida)
    instruccion_idioma = (
        "Redacta toda la minuta en el mismo idioma predominante del material."
        if idioma_salida == "source"
        else f"Redacta TODA la minuta en {etiqueta}, incluyendo titulos, listas, tabla y nota de cautela."
    )
    return f"""
Eres un asistente local que redacta una minuta formal a partir de una transcripcion
o de resumenes intermedios. Usa unicamente el material dado.

Archivo fuente: {nombre_fuente}
Idioma de salida: {etiqueta}
{instruccion_idioma}

{REGLAS_PROMPT}

Devuelve EXACTAMENTE esta estructura en Markdown, sin secciones extra:

# Minuta de reunión

## Archivo fuente
{nombre_fuente}

## Resumen ejecutivo
Resumen breve de lo hablado.

## Temas tratados
- ...

## Acuerdos
- ...

## Tareas y responsables
| Responsable | Tarea | Fecha límite |
|---|---|---|
| No especificado | ... | No especificado |

## Puntos pendientes
- ...

## Riesgos, dudas o temas por aclarar
- ...

## Próximos pasos sugeridos
- ...

## Nota de cautela
Esta minuta fue generada automáticamente a partir de una transcripción. Debe ser revisada antes de usarse como documento oficial.

Material disponible:
\"\"\"
{material_texto}
\"\"\"
""".strip()


def generar_minuta_de_texto(
    texto_limpio: str,
    nombre_fuente: str,
    config: AppConfig,
    callbacks: ProcessCallbacks,
    idioma_salida: str,
) -> str:
    fragmentos = dividir_texto_largo(texto_limpio, config.tamano_fragmento_minuta)
    if not fragmentos:
        raise ValueError("El texto esta vacio; no se puede generar minuta.")

    if len(fragmentos) == 1:
        callbacks.file_progress(20)
        minuta = llamar_a_ollama(crear_prompt_minuta_final(fragmentos[0], nombre_fuente, idioma_salida), config)
        callbacks.file_progress(95)
        return minuta

    resumenes = []
    for i, fragmento in enumerate(fragmentos, start=1):
        callbacks.status(f"Resumiendo fragmento {i} de {len(fragmentos)}...")
        callbacks.file_progress(int((i - 1) / len(fragmentos) * 70))
        prompt = crear_prompt_resumen_fragmento(fragmento, i, len(fragmentos), nombre_fuente, idioma_salida)
        respuesta = llamar_a_ollama(prompt, config)
        resumenes.append(f"## Fragmento {i}\n{respuesta}")

    callbacks.status("Consolidando minuta final...")
    callbacks.file_progress(80)
    minuta = llamar_a_ollama(crear_prompt_minuta_final(resumenes, nombre_fuente, idioma_salida), config)
    callbacks.file_progress(95)
    return minuta


def nombre_base_minuta(base: str, idioma_salida: str = "source", incluir_sufijo: bool = False) -> str:
    base = nombre_seguro(base)
    if incluir_sufijo:
        sufijo = sufijo_idioma_minuta(idioma_salida)
        if sufijo:
            return f"{base}_MINUTA_{sufijo}"
    return f"{base}_MINUTA"


def guardar_minuta_txt(
    minuta_markdown: str,
    carpeta: Path,
    base: str,
    sobrescribir: bool,
    idioma_salida: str = "source",
    incluir_sufijo: bool = False,
) -> Path:
    carpeta.mkdir(parents=True, exist_ok=True)
    ruta = ruta_salida_unica(
        carpeta / f"{nombre_base_minuta(base, idioma_salida, incluir_sufijo)}.txt",
        sobrescribir,
    )
    ruta.write_text(minuta_markdown.strip() + "\n", encoding="utf-8")
    return ruta


def _es_linea_tabla(linea: str) -> bool:
    linea = linea.strip()
    return linea.startswith("|") and linea.endswith("|") and linea.count("|") >= 2


def _parsear_fila_tabla(linea: str) -> List[str]:
    return [celda.strip() for celda in linea.strip().strip("|").split("|")]


def _es_separador_markdown(linea: str) -> bool:
    celdas = _parsear_fila_tabla(linea)
    return bool(celdas) and all(re.fullmatch(r":?-{3,}:?", c.strip()) for c in celdas)


def _agregar_tabla_markdown(doc: Document, lineas: Sequence[str]) -> bool:
    filas = [_parsear_fila_tabla(l) for l in lineas if _es_linea_tabla(l)]
    if len(filas) < 2:
        return False
    if len(filas) >= 2 and _es_separador_markdown(lineas[1]):
        filas = [filas[0]] + filas[2:]
    if not filas:
        return False
    cols = max(len(f) for f in filas)
    tabla = doc.add_table(rows=0, cols=cols)
    tabla.style = "Table Grid"
    tabla.alignment = WD_TABLE_ALIGNMENT.CENTER
    for fila in filas:
        celdas = tabla.add_row().cells
        for i in range(cols):
            celdas[i].text = fila[i] if i < len(fila) else ""
    return True


def _agregar_markdown_basico_a_docx(doc: Document, markdown: str) -> None:
    lineas = markdown.splitlines()
    i = 0
    while i < len(lineas):
        linea = lineas[i].rstrip()
        limpia = linea.strip()
        if not limpia:
            i += 1
            continue
        if _es_linea_tabla(limpia):
            bloque = []
            while i < len(lineas) and _es_linea_tabla(lineas[i].strip()):
                bloque.append(lineas[i].strip())
                i += 1
            if not _agregar_tabla_markdown(doc, bloque):
                for l in bloque:
                    doc.add_paragraph(l)
            continue
        if limpia.startswith("# "):
            doc.add_heading(limpia[2:].strip(), level=0)
        elif limpia.startswith("## "):
            doc.add_heading(limpia[3:].strip(), level=1)
        elif limpia.startswith("### "):
            doc.add_heading(limpia[4:].strip(), level=2)
        elif limpia.startswith("- "):
            doc.add_paragraph(limpia[2:].strip(), style="List Bullet")
        else:
            doc.add_paragraph(limpia)
        i += 1


def guardar_minuta_docx(
    minuta_markdown: str,
    carpeta: Path,
    base: str,
    sobrescribir: bool,
    idioma_salida: str = "source",
    incluir_sufijo: bool = False,
) -> Path:
    carpeta.mkdir(parents=True, exist_ok=True)
    ruta = ruta_salida_unica(
        carpeta / f"{nombre_base_minuta(base, idioma_salida, incluir_sufijo)}.docx",
        sobrescribir,
    )
    doc = Document()
    _agregar_markdown_basico_a_docx(doc, minuta_markdown)
    doc.save(ruta)
    return ruta


# =============================================================================
# PROCESADOR
# =============================================================================

class GamboaProcessor:
    def __init__(self, config: AppConfig, callbacks: Optional[ProcessCallbacks] = None):
        self.config = config
        self.callbacks = callbacks or ProcessCallbacks()
        self.model = None
        self.device_used: Optional[str] = None
        self.compute_used: Optional[str] = None

    def _log(self, msg: str) -> None:
        self.callbacks.log(log_line(msg))

    def cargar_modelo_whisper(self) -> None:
        from faster_whisper import WhisperModel

        if self.model is not None:
            return

        mode = self.config.device_mode.lower()
        wants_cuda = mode in {"auto", "cuda"}

        if wants_cuda:
            self._log(
                f"Cargando Whisper '{self.config.whisper_model_name}' en GPU "
                f"({self.config.gpu_device}/{self.config.gpu_compute_type})..."
            )
            try:
                self.model = WhisperModel(
                    self.config.whisper_model_name,
                    device=self.config.gpu_device,
                    compute_type=self.config.gpu_compute_type,
                )
                self.device_used = self.config.gpu_device
                self.compute_used = self.config.gpu_compute_type
                self._log("Modelo cargado en GPU CUDA.")
                return
            except Exception as exc:
                self._log(f"CUDA no disponible o fallo al cargar. Se usara CPU. Detalle: {exc}")
                if mode == "cuda":
                    self._log("Aunque se eligio CUDA, se continua en CPU para no detener el proceso.")

        self._log(
            f"Cargando Whisper en CPU ({self.config.cpu_compute_type}, "
            f"{self.config.cpu_threads} hilos)..."
        )
        self.model = WhisperModel(
            self.config.whisper_model_name,
            device=self.config.cpu_device,
            compute_type=self.config.cpu_compute_type,
            cpu_threads=self.config.cpu_threads,
        )
        self.device_used = self.config.cpu_device
        self.compute_used = self.config.cpu_compute_type
        self._log("Modelo cargado en CPU.")

    def liberar_modelo_whisper(self) -> None:
        if self.model is None:
            return
        self._log("Liberando modelo Whisper antes de generar minutas...")
        self.model = None
        try:
            gc.collect()
        except Exception:
            pass
        if self.config.limpieza_cuda_agresiva:
            try:
                import torch

                if torch.cuda.is_available():
                    torch.cuda.empty_cache()
                    torch.cuda.ipc_collect()
                    self._log("Limpieza CUDA agresiva ejecutada.")
            except Exception as exc:
                self._log(f"No se pudo ejecutar limpieza CUDA agresiva: {exc}")
        self._log("Modelo Whisper liberado. En procesos largos puede ser necesario reiniciar la app para liberar toda la VRAM.")

    def transcribir_archivo(self, archivo: Path) -> Tuple[List[Dict[str, Union[float, str]]], str, object]:
        if self.model is None:
            self.cargar_modelo_whisper()
        assert self.model is not None

        self.callbacks.status("Transcribiendo...")
        segmentos_generador, info = self.model.transcribe(
            str(archivo),
            language=self.config.language,
            task=self.config.task,
            beam_size=self.config.beam_size,
            vad_filter=self.config.vad_filter,
        )
        duracion = float(getattr(info, "duration", 0.0) or 0.0)
        idioma = getattr(info, "language", None) or "No especificado"
        probabilidad = getattr(info, "language_probability", None)
        if probabilidad is not None:
            self._log(f"Idioma detectado por Whisper: {idioma} ({float(probabilidad):.2%})")
        else:
            self._log(f"Idioma detectado por Whisper: {idioma}")

        segmentos: List[Dict[str, Union[float, str]]] = []
        for segmento in segmentos_generador:
            texto = (segmento.text or "").strip()
            if not texto:
                continue
            item = {"start": float(segmento.start), "end": float(segmento.end), "text": texto}
            segmentos.append(item)
            if duracion > 0:
                percent = min(99, max(1, int(float(segmento.end) / duracion * 100)))
                self.callbacks.file_progress(percent)
        self.callbacks.file_progress(100)
        texto_limpio = " ".join(str(s["text"]).strip() for s in segmentos).strip()
        return segmentos, texto_limpio, info

    def _guardar_log_archivo(self, carpeta: Path, base: str, lineas: Sequence[str]) -> Optional[Path]:
        try:
            ruta = ruta_salida_unica(carpeta / f"{base}_log.txt", self.config.sobrescribir_salidas)
            ruta.write_text("\n".join(lineas).strip() + "\n", encoding="utf-8")
            return ruta
        except Exception as exc:
            self._log(f"No se pudo guardar log de archivo: {exc}")
            return None

    def procesar_audio_video(self, archivo: Path) -> ProcessResult:
        inicio = time.time()
        base = nombre_seguro(archivo.stem)
        carpeta = output_dir_para_archivo(archivo, self.config)
        lineas_log: List[str] = []

        def local_log(msg: str) -> None:
            linea = log_line(msg)
            lineas_log.append(linea)
            self.callbacks.log(linea)

        result = ProcessResult(
            source=str(archivo),
            output_dir=str(carpeta),
            dispositivo=self.device_used,
            compute_type=self.compute_used,
            idioma_configurado=etiqueta_idioma_configurado(self.config.language),
        )
        try:
            self.callbacks.current_file(str(archivo))
            self.callbacks.file_progress(0)
            local_log(f"Transcribiendo archivo: {archivo.name}")
            segmentos, texto_limpio, info = self.transcribir_archivo(archivo)
            idioma_detectado = getattr(info, "language", None) or "No especificado"
            probabilidad = getattr(info, "language_probability", None)
            probabilidad_float = float(probabilidad) if probabilidad is not None else None
            result.idioma_detectado = idioma_detectado
            result.probabilidad_idioma = probabilidad_float

            opts = self.config.output_options
            outputs: Dict[str, str] = {}
            if opts.txt_con_tiempos:
                local_log("Guardando TXT con tiempos...")
                ruta = guardar_txt_con_tiempos(
                    segmentos,
                    carpeta / f"{base}_con_tiempos.txt",
                    self.config.sobrescribir_salidas,
                )
                outputs["txt_con_tiempos"] = str(ruta)
            if opts.txt_limpio:
                local_log("Guardando TXT limpio...")
                ruta = guardar_txt_limpio(
                    texto_limpio,
                    carpeta / f"{base}_texto_limpio.txt",
                    self.config.sobrescribir_salidas,
                )
                outputs["txt_limpio"] = str(ruta)
            else:
                outputs["_texto_limpio_memoria"] = texto_limpio
            if opts.srt:
                local_log("Guardando SRT...")
                ruta = guardar_srt(segmentos, carpeta / f"{base}.srt", self.config.sobrescribir_salidas)
                outputs["srt"] = str(ruta)
            if opts.docx_transcripcion:
                local_log("Guardando DOCX de transcripcion...")
                ruta = guardar_docx_transcripcion(
                    archivo,
                    segmentos,
                    texto_limpio,
                    carpeta / f"{base}.docx",
                    self.config.whisper_model_name,
                    etiqueta_idioma_configurado(self.config.language),
                    idioma_detectado,
                    probabilidad_float,
                    self.device_used or "No especificado",
                    self.compute_used or "No especificado",
                    self.config.sobrescribir_salidas,
                )
                outputs["docx_transcripcion"] = str(ruta)

            result.ok = True
            result.outputs = outputs
            result.dispositivo = self.device_used
            result.compute_type = self.compute_used
            local_log("Transcripcion guardada correctamente.")
        except Exception as exc:
            result.error = str(exc)
            local_log(f"ERROR en archivo {archivo.name}: {exc}")
            self.callbacks.error(f"{archivo.name}: {exc}")
            lineas_log.append(traceback.format_exc())
        finally:
            result.elapsed_seconds = time.time() - inicio
            log_path = self._guardar_log_archivo(carpeta, base, lineas_log)
            if log_path:
                result.outputs["log"] = str(log_path)
        return result

    def generar_minuta_para_resultado(self, result: ProcessResult) -> ProcessResult:
        if not result.ok:
            return result

        archivo = Path(result.source)
        base = nombre_seguro(archivo.stem)
        carpeta = Path(result.output_dir or output_dir_para_archivo(archivo, self.config))

        self.callbacks.current_file(str(archivo))
        self.callbacks.status("Generando minuta...")
        self.callbacks.file_progress(0)
        self._log(f"Generando minuta para: {archivo.name}")

        try:
            if "txt_limpio" in result.outputs:
                texto = Path(result.outputs["txt_limpio"]).read_text(encoding="utf-8", errors="replace")
            else:
                texto = result.outputs.get("_texto_limpio_memoria", "")

            idiomas = resolver_idiomas_minuta(self.config, result.idioma_detectado)
            multi_idioma = len(idiomas) > 1
            for idioma in idiomas:
                etiqueta = etiqueta_idioma_minuta(idioma)
                self.callbacks.status(f"Generando minuta en {etiqueta}...")
                self._log(f"Generando minuta en {etiqueta} para: {archivo.name}")
                minuta = generar_minuta_de_texto(texto, archivo.name, self.config, self.callbacks, idioma)
                ruta_txt = guardar_minuta_txt(
                    minuta,
                    carpeta,
                    base,
                    self.config.sobrescribir_salidas,
                    idioma_salida=idioma,
                    incluir_sufijo=multi_idioma,
                )
                key_suffix = sufijo_idioma_minuta(idioma) if multi_idioma else ""
                result.outputs[f"minuta_txt{('_' + key_suffix) if key_suffix else ''}"] = str(ruta_txt)
                if self.config.output_options.docx_minuta:
                    ruta_docx = guardar_minuta_docx(
                        minuta,
                        carpeta,
                        base,
                        self.config.sobrescribir_salidas,
                        idioma_salida=idioma,
                        incluir_sufijo=multi_idioma,
                    )
                    result.outputs[f"minuta_docx{('_' + key_suffix) if key_suffix else ''}"] = str(ruta_docx)
                result.idiomas_minuta.append(idioma)
            result.minuta_generada = True
            self.callbacks.file_progress(100)
            self._log(f"Minuta generada: {archivo.name}")
        except Exception as exc:
            result.error = f"{result.error or ''} Error minuta: {exc}".strip()
            self.callbacks.error(f"Minuta {archivo.name}: {exc}")
            self._log(f"ERROR generando minuta para {archivo.name}: {exc}")
        return result

    def generar_minuta_desde_txt(self, txt_path: Path) -> ProcessResult:
        inicio = time.time()
        base = base_desde_txt(txt_path)
        carpeta = output_dir_para_txt(txt_path, self.config)
        result = ProcessResult(source=str(txt_path), output_dir=str(carpeta))

        try:
            self.callbacks.current_file(str(txt_path))
            self.callbacks.status("Leyendo TXT...")
            texto = txt_path.read_text(encoding="utf-8", errors="replace")
            self.callbacks.status("Generando minuta...")
            idiomas = resolver_idiomas_minuta(self.config, None)
            multi_idioma = len(idiomas) > 1
            for idioma in idiomas:
                etiqueta = etiqueta_idioma_minuta(idioma)
                self.callbacks.status(f"Generando minuta en {etiqueta}...")
                minuta = generar_minuta_de_texto(texto, txt_path.name, self.config, self.callbacks, idioma)
                ruta_txt = guardar_minuta_txt(
                    minuta,
                    carpeta,
                    base,
                    self.config.sobrescribir_salidas,
                    idioma_salida=idioma,
                    incluir_sufijo=multi_idioma,
                )
                key_suffix = sufijo_idioma_minuta(idioma) if multi_idioma else ""
                result.outputs[f"minuta_txt{('_' + key_suffix) if key_suffix else ''}"] = str(ruta_txt)
                if self.config.output_options.docx_minuta:
                    ruta_docx = guardar_minuta_docx(
                        minuta,
                        carpeta,
                        base,
                        self.config.sobrescribir_salidas,
                        idioma_salida=idioma,
                        incluir_sufijo=multi_idioma,
                    )
                    result.outputs[f"minuta_docx{('_' + key_suffix) if key_suffix else ''}"] = str(ruta_docx)
                result.idiomas_minuta.append(idioma)
            result.ok = True
            result.minuta_generada = True
            self._log(f"Minuta generada desde TXT: {txt_path.name}")
        except Exception as exc:
            result.error = str(exc)
            self.callbacks.error(f"{txt_path.name}: {exc}")
            self._log(f"ERROR generando minuta desde TXT {txt_path.name}: {exc}")
        finally:
            result.elapsed_seconds = time.time() - inicio
        return result

    def process(self, mode: str, inputs: Sequence[Path]) -> ProcessSummary:
        start = time.time()
        results: List[ProcessResult] = []
        cancelled = False
        idiomas_minuta_plan = resolver_idiomas_minuta(self.config, None)
        unidades_minuta_por_archivo = max(1, len(idiomas_minuta_plan))

        if mode == "txt_to_minute":
            self._log("Verificando Ollama...")
            ollama = verificar_ollama(self.config.ollama_model, self.config.ollama_tags_url)
            self._log(str(ollama["message"]))
            if not ollama["available"] or not ollama["model_installed"]:
                raise RuntimeError(str(ollama["message"]))

            total = len(inputs) * unidades_minuta_por_archivo
            for i, txt in enumerate(inputs, start=1):
                self.callbacks.general_progress((i - 1) * unidades_minuta_por_archivo, total)
                results.append(self.generar_minuta_desde_txt(Path(txt)))
                self.callbacks.general_progress(i * unidades_minuta_por_archivo, total)
                if self.callbacks.should_stop_after_current():
                    cancelled = True
                    break
            return ProcessSummary(results, self._summary_output_root(results), time.time() - start, cancelled)

        files: List[Path] = []
        for item in inputs:
            item = Path(item)
            if item.is_dir():
                files.extend(recolectar_archivos_compatibles(item))
            elif item.is_file() and item.suffix.lower() in EXTENSIONES_COMPATIBLES:
                files.append(item)

        if not files:
            raise RuntimeError("No se encontraron archivos compatibles.")

        if mode == "transcribe_minute" and self.config.whisper_model_name == "large-v3":
            self._log("Advertencia: large-v3 puede ocupar bastante VRAM. Si Ollama queda lento, use Solo transcripcion y luego Minuta desde TXT.")

        self.cargar_modelo_whisper()

        total_units = len(files) + (len(files) * unidades_minuta_por_archivo if mode == "transcribe_minute" else 0)
        done_units = 0
        for idx, archivo in enumerate(files, start=1):
            self.callbacks.status(f"Procesando archivo {idx} de {len(files)}")
            self.callbacks.general_progress(done_units, total_units)
            result = self.procesar_audio_video(archivo)
            results.append(result)
            done_units += 1
            self.callbacks.general_progress(done_units, total_units)
            if self.callbacks.should_stop_after_current():
                cancelled = True
                self._log("Detencion solicitada. Se detendra despues del archivo actual.")
                break

        if mode == "transcribe_minute" and not cancelled:
            ollama = verificar_ollama(self.config.ollama_model, self.config.ollama_tags_url)
            self._log(str(ollama["message"]))
            if not ollama["available"] or not ollama["model_installed"]:
                self._log("Ollama no disponible. Las transcripciones quedan guardadas; no se generaran minutas.")
            else:
                if self.config.liberar_whisper_antes_ollama:
                    self.liberar_modelo_whisper()
                for result in results:
                    if result.ok:
                        self.generar_minuta_para_resultado(result)
                    done_units += unidades_minuta_por_archivo
                    self.callbacks.general_progress(done_units, total_units)
                    if self.callbacks.should_stop_after_current():
                        cancelled = True
                        break

        self.callbacks.status("Proceso terminado.")
        self.callbacks.file_progress(100)
        self.callbacks.general_progress(total_units if not cancelled else done_units, total_units)
        return ProcessSummary(results, self._summary_output_root(results), time.time() - start, cancelled)

    def _summary_output_root(self, results: Sequence[ProcessResult]) -> Optional[str]:
        for result in results:
            if result.output_dir:
                return str(Path(result.output_dir).parent)
        if self.config.output_root:
            return str(self.config.output_root)
        return None


def exportar_reporte_ejecucion(summary: ProcessSummary, ruta: Union[str, Path]) -> Path:
    ruta = Path(ruta)
    ruta.parent.mkdir(parents=True, exist_ok=True)
    data = {
        "app": APP_NAME,
        "fecha": now_label(),
        "elapsed_seconds": summary.elapsed_seconds,
        "cancelled": summary.cancelled,
        "ok_count": summary.ok_count,
        "error_count": summary.error_count,
        "minute_count": summary.minute_count,
        "output_root": summary.output_root,
        "results": [r.__dict__ for r in summary.results],
    }
    ruta.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")
    return ruta
